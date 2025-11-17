
"""
doc_redactor.py

Full pipeline:
1. Ask user to select image/pdf/docx/txt
2. OCR via OCR.space (primary) or local pytesseract fallback
3. Clean/format text with Gemini (optional, if GEMINI_API_KEY is set)
4. Redact PII by label + regex patterns
5. Save cleaned & redacted DOCX and PDF

Set env vars:
export GEMINI_API_KEY="..."    # optional; if missing, AI cleaning will be skipped
export OCR_API_KEY="..."       # optional; if missing, local pytesseract is used
"""

import os
import re
import io
import sys
import requests
from tkinter import Tk, filedialog, messagebox
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# Optional local OCR fallback and PDF/image libs
try:
    import fitz  # PyMuPDF (for PDF image extraction)
    from pdf2image import convert_from_path
    from PIL import Image
    import pytesseract
except Exception:
    fitz = None
    convert_from_path = None
    Image = None
    pytesseract = None

# Google generative ai (Gemini) - optional
try:
    import google.generativeai as genai
except Exception:
    genai = None

# ------------------------------
# CONFIGURATION / KEYS
# ------------------------------
OCR_API_KEY = os.getenv("OCR_API_KEY", "K85450490888957").strip()  # optional
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyClP2B1jdADvbxd8I96w5Fok8aZZQfXEbQ").strip()  # optional

if genai and GEMINI_API_KEY:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
    except Exception as e:
        print(f"[WARN] Could not configure Gemini: {e}")
        genai = None
else:
    # If library not available or key missing, we won't call Gemini
    genai = None

# Choose Gemini model if genai is available
GEMINI_MODEL = "gemini-1.5-pro-latest"

# ------------------------------
# PII LABELS & PATTERNS
# ------------------------------
PII_LABELS = [

    # -----------------------------
    # 1. Government Issued ID
    # -----------------------------
    "government issued id", "Government Issued ID", "GOVERNMENT ISSUED ID",
    "govt issued id", "gov issued id", "gov issued identification",
    "gov id", "govt id", "government id", "government identification",
    "id issued by government", "government identity card",
    "id card", "identity card", "identification id",
    "official id", "official identification", "national id",
    "national identification", "gov identity",

    # -----------------------------
    # 2. Social Security Number
    # -----------------------------
    "social security number", "Social Security Number", "SOCIAL SECURITY NUMBER",
    "ssn", "SSN", "S.S.N.", "social security no", "ss number",
    "soc sec no", "ssn number", "social sec number", "social security #",

    # -----------------------------
    # 3. Tax ID
    # -----------------------------
    "tax id", "Tax ID", "TAX ID", "tax identification number",
    "tin", "TIN", "T.I.N.", "tax no", "tax number",
    "taxpayer id", "tax payer number",

    # -----------------------------
    # 4. Federal Employer ID
    # -----------------------------
    "federal employer id", "Federal Employer ID", "FEDERAL EMPLOYER ID",
    "employer id", "employer identification", "feid", "FEID", "F.E.I.D.",

    # -----------------------------
    # 5. FEIN
    # -----------------------------
    "fein", "FEIN", "F.E.I.N.", "federal employer identification number",
    "fein number", "federal ein", "employer ein",

    # -----------------------------
    # 6. Driver's License
    # -----------------------------
    "driver's license", "Driver's License", "Driver' s License","License","DRIVER'S LICENSE",
    "drivers license", "driver license", "driving license",
    "dl number", "DL", "D.L.", "license number", "driver id",

    # -----------------------------
    # 7. Identification Card
    # -----------------------------
    "identification card", "Identification Card", "ID card",
    "identity card", "id", "ID", "identification", "id number",
    "identification number",

    # -----------------------------
    # 8. Passport
    # -----------------------------
    "passport", "Passport", "PASSPORT", "passport number",
    "passport no", "pp number", "passport id",

    # -----------------------------
    # 9. Military ID
    # -----------------------------
    "military id", "Military ID", "MILITARY ID",
    "army id", "navy id", "airforce id", "defense id",
    "military identification",

    # -----------------------------
    # 10. Date of Birth
    # -----------------------------
    "date of birth", "Date of Birth", "DATE OF BIRTH",
    "dob", "DOB", "birth date", "birth info","D.o.B.","DOB",
    "date born", "born on", "birthdate","D.O.B.",

    # -----------------------------
    # 11. Home Address
    # -----------------------------
    "home address", "Home Address", "HOME ADDRESS",
    "residential address", "residence address", "address", "addr","ADDRESS",
    "street address", "street addr", "residential addr","Address",

    # -----------------------------
    # 12. Home Telephone Number
    # -----------------------------
    "home telephone number", "Home Telephone number",
    "HOME TELEPHONE NUMBER", "telephone number",
    "home phone", "landline", "tel number",

    # -----------------------------
    # 13. Cell Phone Number
    # -----------------------------
    "cell phone number", "Cell phone number", "CELL PHONE NUMBER",
    "mobile number", "mobile no", "cell number", "phone number",
    "contact number", "contact no","ph number","Cell No",

    # -----------------------------
    # 14. Email Address
    # -----------------------------
    "email address", "Email Address", "EMAIL ADDRESS",
    "email", "e-mail", "email id", "mail id","Email","email ID","eMail","gmail","g-mail",

    # -----------------------------
    # 15. Social Media Contact Information
    # -----------------------------
    "social media contact information", "Social Media Contact Information",
    "SOCIAL MEDIA CONTACT INFORMATION", "social media info",
    "social handle", "social contact", "social media account",

    # -----------------------------
    # 16. Health Insurance Policy Number
    # -----------------------------
    "health insurance policy number", "Health Insurance Policy Number",
    "insurance policy number", "policy number", "policy no",
    "health insurance number", "insurance number",

    # -----------------------------
    # 17. Medical Record Number
    # -----------------------------
    "medical record number", "Medical Record Number",
    "MRN", "mrn", "medical record no", "med record number","medical","record","number",

    # -----------------------------
    # 18. Claim Number
    # -----------------------------
    "claim number", "Claim Number", "CLAIM NUMBER",
    "claim no", "claim id",

    # -----------------------------
    # 19. Patient Account Number
    # -----------------------------
    "patient account number", "Patient Account Number",
    "PATIENT ACCOUNT NUMBER", "patient id", "patient account",

    # -----------------------------
    # 20. File Number
    # -----------------------------
    "file number", "File Number", "FILE NUMBER",
    "file no", "file id", "file reference",

    # -----------------------------
    # 21. Chart Number
    # -----------------------------
    "chart number", "Chart Number", "CHART NUMBER",
    "chart no", "chart id",

    # -----------------------------
    # 22. Individual Financial Account Number
    # -----------------------------
    "individual financial account number", "Individual Financial Account Number",
    "financial account number", "financial account", "account number",

    # -----------------------------
    # 23. Bank Account Number
    # -----------------------------
    "bank account number", "Bank Account Number", "BANK ACCOUNT NUMBER",
    "bank no", "account no", "acct number",

    # -----------------------------
    # 24. Financial Information
    # -----------------------------
    "financial information", "Financial Information",
    "FINANCIAL INFORMATION", "financial data", "financial details",

    # -----------------------------
    # 25. Credit Card Number
    # -----------------------------
    "credit card number", "Credit Card Number", "CREDIT CARD NUMBER",
    "credit card", "card number", "cc number", "card no"
]


PATTERNS = {
    "SSN": r"\b\d{3}-\d{2}-\d{4}\b",
    "Phone": r"\b\+?\d{1,3}?[-.\s]?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}\b",
    "Email": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",
    "Credit Card": r"\b(?:\d{4}[-\s]?){3}\d{4}\b",
    "Date (MM/DD/YYYY)": r"\b(?:0?[1-9]|1[0-2])[\/\-.](?:0?[1-9]|[12]\d|3[01])[\/\-.]\d{4}\b",
    # Add more as needed
}

# ------------------------------
# HELPERS: filetype detection
# ------------------------------
def detect_filetype(path):
    ext = os.path.splitext(path)[1].lower()
    if ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif"]:
        return "image"
    if ext == ".pdf":
        return "pdf"
    if ext == ".txt":
        return "text"
    if ext == ".docx":
        return "docx"
    return None

# ------------------------------
# OCR.space extractor
# ------------------------------
def ocrspace_extract(path, language="eng"):
    """Use OCR.space API to extract text. Returns empty string on failure."""
    if not OCR_API_KEY:
        return ""
    try:
        with open(path, "rb") as f:
            files = {"file": (os.path.basename(path), f)}
            data = {
                "apikey": OCR_API_KEY,
                "language": language,
                "OCREngine": 2,
                "isTable": False,
                "scale": True
            }
            resp = requests.post("https://api.ocr.space/parse/image", files=files, data=data, timeout=120)
            resp.raise_for_status()
            rjson = resp.json()
            if rjson.get("IsErroredOnProcessing"):
                print("[OCR.space] Error on processing:", rjson.get("ErrorMessage"))
                return ""
            parsed_results = rjson.get("ParsedResults")
            if not parsed_results:
                return ""
            # Join parsed results across pages
            return "\n".join(pr.get("ParsedText", "") for pr in parsed_results).strip()
    except Exception as e:
        print("[OCR.space] exception:", e)
        return ""

# ------------------------------
# pytesseract extractor (fallback)
# ------------------------------
def pytesseract_extract(path):
    """Local OCR fallback using pytesseract (image or PDF -> images)."""
    if pytesseract is None or Image is None:
        return ""
    try:
        ftype = detect_filetype(path)
        if ftype == "image":
            img = Image.open(path)
            return pytesseract.image_to_string(img)
        elif ftype == "pdf":
            text_parts = []
            # Prefer pdf2image if available for better rendering
            if convert_from_path:
                images = convert_from_path(path, dpi=200)
                for img in images:
                    text_parts.append(pytesseract.image_to_string(img))
            else:
                # Try PyMuPDF to render pages
                if fitz:
                    doc = fitz.open(path)
                    for page in doc:
                        pix = page.get_pixmap(dpi=150)
                        img = Image.open(io.BytesIO(pix.tobytes()))
                        text_parts.append(pytesseract.image_to_string(img))
                else:
                    return ""
            return "\n".join(text_parts)
    except Exception as e:
        print("[pytesseract] exception:", e)
        return ""

# ------------------------------
# EXTRACT TEXT FROM FILE (ALL TYPES)
# ------------------------------
def extract_text_from_file(path):
    """
    Main unified extractor:
    - TXT -> direct read
    - DOCX -> docx paragraphs
    - PDF/IMAGE -> OCR.space primary -> pytesseract fallback
    Returns extracted text (string) or empty string on failure.
    """
    filetype = detect_filetype(path)
    if not filetype:
        messagebox.showerror("Error", f"Unsupported file type: {path}")
        return ""

    # TXT files
    if filetype == "text":
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            print("[extract_text_from_file] TXT read failed:", e)
            return ""

    # DOCX files
    if filetype == "docx":
        try:
            doc = Document(path)
            # Extract text from paragraphs and tables (some docx contain tables)
            paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            # Extract text from tables too
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            paragraphs.append(cell.text)
            return "\n".join(paragraphs)
        except Exception as e:
            print("[extract_text_from_file] DOCX read failed:", e)
            return ""

    # Image or PDF -> OCR
    if filetype in ("image", "pdf"):
        # Try OCR.space first if key present
        if OCR_API_KEY:
            txt = ocrspace_extract(path)
            if txt and txt.strip():
                return txt
        # Fallback to local OCR
        txt = pytesseract_extract(path)
        if txt and txt.strip():
            return txt
        # Try to extract embedded text from PDF (some PDFs contain text not images)
        if filetype == "pdf":
            try:
                if fitz:
                    doc = fitz.open(path)
                    text_parts = []
                    for page in doc:
                        ptext = page.get_text("text")
                        if ptext and ptext.strip():
                            text_parts.append(ptext)
                    if text_parts:
                        return "\n".join(text_parts)
            except Exception:
                pass
        return ""

    return ""

# ------------------------------
# GEMINI: structure fixing / formatting (optional)
# ------------------------------
def fix_text_structure_with_gemini(text, max_tokens=4096):
    """
    Uses Gemini (if available) to clean and reformat OCR text.
    Returns cleaned text (or original if AI not available or fails).
    """
    if genai is None:
        return text

    try:
        # Use the high-level generate_content / generate API depending on the genai version.
        model = genai.GenerativeModel(GEMINI_MODEL)

        prompt = f"""You are an expert text cleaner. Clean and reflow the OCR-extracted text.
1) Fix spacing, remove false line-breaks inside paragraphs but keep real paragraphs.
2) Correct obvious OCR mis-reads (like 'l' vs '1', 'O' vs '0' when clearly wrong).
3) Preserve headings and lists.
4) Return ONLY the cleaned text, no extra commentary.

OCR_TEXT_START:
{text}
OCR_TEXT_END:
"""

        # Some SDKs allow .generate_content or .generate — try generate_content first
        cleaned = ""
        try:
            response = model.generate_content([prompt])
            # response may have 'candidates' or 'output' fields
            if hasattr(response, "text") and response.text:
                cleaned = response.text
            elif isinstance(response, (list, tuple)) and len(response) > 0:
                first = response[0]
                if isinstance(first, dict):
                    cleaned = first.get("content") or first.get("text") or first.get("output", "")
                else:
                    cleaned = str(first)
            elif isinstance(response, dict):
                cleaned = response.get("output", "") or response.get("content", "") or response.get("text", "")
            else:
                cleaned = str(response)
        except Exception:
            # try another common method name
            response2 = model.generate(prompt=prompt)
            if response2:
                # try parse
                if isinstance(response2, dict):
                    cleaned = response2.get("candidates", [{}])[0].get("content") or response2.get("output", "")
                else:
                    cleaned = str(response2)

        if not cleaned:
            return text
        return cleaned.strip()
    except Exception as e:
        print("[Gemini] AI formatting failed:", e)
        return text

# ------------------------------
# REDACTION helpers
# ------------------------------
def blackout_text(s):
    if not s:
        return s
    # keep same length blackout for visual parity
    return "█" * len(s)

def redact_label_values(text, labels):
    """Redact values that appear after known labels like 'Email: <value>'."""
    redacted = text
    for label in labels:
        label_escaped = re.escape(label)
        patterns = [
            rf"({label_escaped}\s*[:\-–]\s*)([^\n\r]+)",   # Label: value
            rf"({label_escaped}\s+)([^\n\r]+)"            # Label value
        ]
        for p in patterns:
            redacted = re.sub(p, lambda m: m.group(1) + blackout_text(m.group(2).strip()), redacted, flags=re.IGNORECASE)
    return redacted

def redact_patterns(text, patterns_dict):
    redacted = text
    for name, patt in patterns_dict.items():
        try:
            redacted = re.sub(patt, lambda m: blackout_text(m.group(0)), redacted, flags=re.IGNORECASE)
        except re.error as e:
            print(f"[redact_patterns] invalid regex for {name}: {e}")
    return redacted

def redact_text_content(text):
    r = redact_label_values(text, PII_LABELS)
    r = redact_patterns(r, PATTERNS)
    return r

# ------------------------------
# SAVE to DOCX and PDF
# ------------------------------
def save_to_word(text, input_path):
    try:
        folder = os.path.dirname(input_path) or os.getcwd()
        filename = os.path.splitext(os.path.basename(input_path))[0]
        out_path = os.path.join(folder, f"{filename}_cleaned_redacted.docx")
        doc = Document()
        # Add a little top heading
        doc.add_heading("Cleaned & Redacted Document", level=2)
        for para in text.splitlines():
            p = para.strip()
            if p:
                doc.add_paragraph(p)
            else:
                # add an empty paragraph to preserve paragraph breaks
                doc.add_paragraph('')
        doc.save(out_path)
        return out_path
    except Exception as e:
        messagebox.showerror("Save Error", f"Failed to save Word document: {e}")
        return None

def save_to_pdf(text, input_path):
    try:
        folder = os.path.dirname(input_path) or os.getcwd()
        filename = os.path.splitext(os.path.basename(input_path))[0]
        out_path = os.path.join(folder, f"{filename}_cleaned_redacted.pdf")
        doc = SimpleDocTemplate(out_path)
        styles = getSampleStyleSheet()
        story = []
        for para in text.splitlines():
            p = para.strip()
            if p:
                story.append(Paragraph(p, styles["Normal"]))
            else:
                # add a blank paragraph to preserve breaks
                story.append(Paragraph(" ", styles["Normal"]))
        doc.build(story)
        return out_path
    except Exception as e:
        messagebox.showerror("PDF Error", f"Failed to create PDF: {e}")
        return None

# ------------------------------
# MAIN PROCESSING PIPELINE
# ------------------------------
def process_file_pipeline():
    root = Tk()
    root.withdraw()
    messagebox.showinfo("Select File", "Choose an image, PDF, DOCX, or TXT file to process and redact.")
    filepath = filedialog.askopenfilename(
        title="Select File",
        filetypes=[
            ("Supported files", "*.jpg;*.jpeg;*.png;*.bmp;*.tiff;*.pdf;*.txt;*.docx"),
            ("All files", "*.*")
        ]
    )
    if not filepath:
        return

    if not os.path.exists(filepath):
        messagebox.showerror("Error", "Selected file does not exist.")
        return

    # Step 1: Extract text
    messagebox.showinfo("Processing", "Extracting text from your file...")
    raw_text = extract_text_from_file(filepath)
    if not raw_text or not raw_text.strip():
        messagebox.showerror("OCR Error", "No text could be extracted. Make sure the file contains readable text or install pytesseract/local OCR.")
        return

    print(f"[INFO] Raw extracted text length: {len(raw_text)}")

    # Step 2: AI cleanup (optional)
    cleaned = raw_text
    if genai is not None:
        messagebox.showinfo("Formatting", f"Cleaning and formatting text with {GEMINI_MODEL} ...")
        cleaned = fix_text_structure_with_gemini(raw_text)
        if not cleaned or not cleaned.strip():
            messagebox.showwarning("AI Warning", "AI returned no cleaned text; using original extracted text.")
            cleaned = raw_text
    else:
        # No Gemini available; proceed with OCR text
        cleaned = raw_text

    print(f"[INFO] Cleaned text length: {len(cleaned)}")

    # Step 3: Redaction
    messagebox.showinfo("Redaction", "Applying redaction to remove PII...")
    redacted = redact_text_content(cleaned)
    print(f"[INFO] Redacted text length: {len(redacted)}")

    # Step 4: Save outputs
    messagebox.showinfo("Saving", "Saving DOCX and PDF outputs...")
    word_path = save_to_word(redacted, filepath)
    pdf_path = save_to_pdf(redacted, filepath)

    success_msg = f"Processing complete!\n\nWord: {word_path or 'Failed'}\nPDF: {pdf_path or 'Failed'}"
    messagebox.showinfo("Done", success_msg)

    if word_path:
        print(f"Saved: {word_path}")
    if pdf_path:
        print(f"Saved: {pdf_path}")

# ------------------------------
# REDACT EXISTING FILE (text/docx) - kept for convenience
# ------------------------------
def redact_existing_file():
    root = Tk()
    root.withdraw()
    filepath = filedialog.askopenfilename(title="Select file to redact", filetypes=[("Text and DOCX", "*.txt;*.docx;*.pdf")])
    if not filepath:
        return
    try:
        content = ""
        ftype = detect_filetype(filepath)
        if ftype == "text":
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
        elif ftype == "docx":
            doc = Document(filepath)
            content = "\n".join(p.text for p in doc.paragraphs)
        elif ftype == "pdf":
            # Try to extract text from PDF (non-scanned)
            if fitz:
                doc = fitz.open(filepath)
                pages = [p.get_text("text") for p in doc]
                content = "\n".join(p for p in pages if p and p.strip())
            else:
                # fallback to OCR-based pipeline
                content = extract_text_from_file(filepath)
        else:
            content = extract_text_from_file(filepath)

        if not content:
            messagebox.showwarning("Empty", "No text found to redact.")
            return

        redacted = redact_text_content(content)
        folder = os.path.dirname(filepath) or os.getcwd()
        filename = os.path.splitext(os.path.basename(filepath))[0]
        out_path = os.path.join(folder, f"{filename}_redacted.txt")
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(redacted)
        messagebox.showinfo("Success", f"Redacted file saved: {out_path}")
    except Exception as e:
        messagebox.showerror("Error", f"Redaction failed: {e}")

# ------------------------------
# MAIN MENU
# ------------------------------
def main():
    root = Tk()
    root.withdraw()
    choice = messagebox.askyesno(
        "Document Processor",
        "YES: Full processing (OCR + optional AI + Redaction)\nNO: Redact existing text/docx/pdf file only"
    )
    if choice:
        process_file_pipeline()
    else:
        redact_existing_file()

if __name__ == "__main__":
    main()
